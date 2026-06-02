"""Parse Word, PDF, or plain text into structured presentation document data."""

from __future__ import annotations

import re
from pathlib import Path
from typing import Any

import pdfplumber
from docx import Document

def truncate(text: str, max_len: int) -> str:
    text = re.sub(r"\s+", " ", text.strip())
    if len(text) <= max_len:
        return text
    cut = text[: max_len - 1].rsplit(" ", 1)[0]
    return cut + "…"


def _split_paragraphs(text: str) -> list[str]:
    chunks = [p.strip() for p in re.split(r"\n\s*\n+", text) if p.strip()]
    if len(chunks) >= 3:
        return chunks
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    return lines if lines else chunks


def paragraphs_to_document(
    paragraphs: list[str],
    *,
    timeline: list[list[str]] | None = None,
) -> dict[str, Any]:
    """Build the deck content dict used by deck_builder (Vienna-style structure)."""
    if not paragraphs:
        return _empty_document()

    sections: list[dict] = []
    current: dict | None = None
    for text in paragraphs[2:]:
        if re.match(r"^\d+\.\s", text) and len(text) < 120:
            if current:
                sections.append(current)
            current = {"heading": text, "paragraphs": [], "pairs": [], "bullets": []}
            continue
        if current is None:
            continue
        if "\n" in text and len(text.split("\n", 1)[0]) < 90:
            parts = text.split("\n", 1)
            current["pairs"].append((parts[0].strip(), parts[1].strip() if len(parts) > 1 else ""))
        elif current["pairs"] and not current["pairs"][-1][1]:
            t, _ = current["pairs"][-1]
            current["pairs"][-1] = (t, text)
        elif len(text) < 90 and text[0].isupper() and not text.endswith("."):
            current["pairs"].append((text, ""))
        elif re.match(r"^\d+\.\s", text) and len(text) < 100:
            current["bullets"].append(text)
        elif text.startswith(("•", "-", "Responsible", "Accountable", "Consulted", "Informed")):
            current["bullets"].append(text)
        else:
            current["paragraphs"].append(text)

    if current:
        sections.append(current)

    if not sections and len(paragraphs) > 2:
        body = paragraphs[2:]
        mid = max(1, len(body) // 3)
        sections = [
            {"heading": "1. What?", "paragraphs": body[:mid], "pairs": [], "bullets": []},
            {"heading": "2. So What?", "paragraphs": body[mid : 2 * mid], "pairs": [], "bullets": []},
            {"heading": "3. Now What?", "paragraphs": body[2 * mid :], "pairs": [], "bullets": []},
        ]

    what = next((s for s in sections if "What" in s["heading"]), {"paragraphs": paragraphs[2:], "pairs": []})
    so_what = next((s for s in sections if "So What" in s["heading"]), {"paragraphs": [], "pairs": []})
    now_what = next((s for s in sections if "Now What" in s["heading"]), {})

    opportunities = list(so_what.get("pairs", []))
    if len(opportunities) < 6:
        opportunities = []
        paras = so_what.get("paragraphs", [])
        i = 0
        while i < len(paras):
            p = paras[i]
            if "focus include" in p.lower() or "discovery phase will" in p.lower():
                i += 1
                continue
            if i + 1 < len(paras) and len(p) < 85:
                opportunities.append((p, paras[i + 1]))
                i += 2
            else:
                i += 1

    workstreams: list[tuple[str, str]] = []
    ws_start: int | None = None
    for i, text in enumerate(paragraphs):
        if "five workstreams" in text.lower() or "workstreams:" in text.lower():
            ws_start = i + 1
            break
    if ws_start is not None:
        for text in paragraphs[ws_start:]:
            if text == "Required Project Team":
                break
            if "\n" in text:
                title, body = text.split("\n", 1)
                workstreams.append((title.strip(), body.strip()))
            if len(workstreams) >= 5:
                break

    team_roles: list[tuple[str, str]] = []
    for text in paragraphs:
        if not re.search(r"\(\s*\d+(\.\d+)?\s*(FTE|fTE)?\s*\)", text):
            continue
        if re.match(r"^Core roles", text, re.I):
            continue
        if "—" in text or "–" in text:
            parts = re.split(r"\s*[—–]\s*", text, maxsplit=1)
            if len(parts) == 2:
                team_roles.append((parts[0].strip(), parts[1].strip()))
        elif re.search(r"\)\s*[-–—]?\s*\S", text):
            m = re.match(r"^(.+?\))\s*[-–—]?\s*(.+)$", text)
            if m:
                team_roles.append((m.group(1).strip(), m.group(2).strip()))

    outputs: list[str] = []
    capture = False
    closing = (
        "This discovery phase will create the fact base required to move from assessment to execution."
    )
    for text in paragraphs:
        if text.startswith("By the end of the discovery") or text.startswith("Expected Outputs"):
            capture = True
            continue
        if capture:
            if "fact base" in text.lower() or text.startswith("Industry Standards"):
                if "fact base" in text.lower():
                    closing = text
                break
            if len(text) > 20 and not text.startswith("1."):
                outputs.append(text)

    placement_options: list[tuple[str, str]] = []
    for text in paragraphs:
        if text.startswith("Remain local"):
            placement_options.append(("Remain local (JV)", text))
        elif text.startswith("Transition to the Prague") or text.startswith("Transition to"):
            placement_options.append(("Transition to SSC", text))
        elif text.startswith("Scale from an existing JV") or text.startswith("Scale from"):
            placement_options.append(("Scale JV capability", text))
        elif text.startswith("Build new centralized") or text.startswith("Build new central"):
            placement_options.append(("Build central capability", text))

    methodology_groups: list[tuple[str, list[str]]] = []
    method_titles = [
        ("Service catalogue & taxonomy", "1. Service Catalogue"),
        ("SIPOC & process mapping", "2. SIPOC"),
        ("Lean Six Sigma diagnostics", "3. Lean"),
        ("RACI & operating model", "4. RACI"),
        ("GBS maturity assessment", "5. Shared Services"),
        ("Fit-gap & service placement", "6. Service Placement"),
        ("Activity-based costing", "7. Activity-Based"),
        ("Benchmarking & prioritization", "8. Benchmarking"),
    ]
    for label, marker in method_titles:
        bullets: list[str] = []
        started = False
        for text in paragraphs:
            if marker in text:
                started = True
                if len(text) < 120:
                    bullets.append(truncate(text, 100))
                continue
            if started:
                if re.match(r"^\d+\.\s", text) and marker not in text:
                    break
                if len(text) < 95 and (
                    text[0].islower() or text.startswith(("Service ", "Process ", "Degree"))
                ):
                    bullets.append(truncate(text, 88))
                elif text.endswith(".") and len(bullets) < 5:
                    bullets.append(truncate(text, 88))
                if len(bullets) >= 4:
                    break
        methodology_groups.append((label, bullets[:4]))

    fit_gap_criteria = [
        t
        for t in paragraphs
        if t
        in (
            "Cost efficiency",
            "Process complexity",
            "Scalability",
            "Regulatory or market-specific requirements",
            "Language or local knowledge dependency",
            "Technology readiness",
            "Data quality and reporting requirements",
            "Required proximity to the customer, merchant, regulator or local management",
            "Existing SSC capability",
            "Existing JV best practice capability",
            "Risk and control implications",
        )
    ]

    prioritization = [
        t
        for t in paragraphs
        if t
        in (
            "Financial benefit",
            "Ease of implementation",
            "Process standardization potential",
            "Risk reduction",
            "Customer or stakeholder impact",
            "Technology dependency",
            "Regulatory complexity",
            "Speed to implement",
            "Readiness for transition",
        )
    ]

    return {
        "title": paragraphs[0],
        "subtitle": paragraphs[1] if len(paragraphs) > 1 else "",
        "sections": sections,
        "what": what,
        "so_what": so_what,
        "now_what": now_what,
        "opportunities": opportunities[:7],
        "workstreams": workstreams[:5],
        "team_roles": team_roles[:9],
        "timeline": timeline or [],
        "outputs": outputs[:10],
        "closing": closing,
        "placement_options": placement_options,
        "methodology_groups": methodology_groups,
        "fit_gap_criteria": fit_gap_criteria,
        "prioritization": prioritization,
    }


def _empty_document() -> dict[str, Any]:
    return {
        "title": "Presentation",
        "subtitle": "Generated with Presentation Studio",
        "sections": [],
        "what": {"paragraphs": ["Add source content to generate slides."], "pairs": []},
        "so_what": {"paragraphs": [], "pairs": []},
        "now_what": {},
        "opportunities": [],
        "workstreams": [],
        "team_roles": [],
        "timeline": [],
        "outputs": [],
        "closing": "Thank you.",
        "placement_options": [],
        "methodology_groups": [],
        "fit_gap_criteria": [],
        "prioritization": [],
    }


def load_docx(path: Path) -> dict[str, Any]:
    doc = Document(str(path))
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    timeline: list[list[str]] = []
    if doc.tables:
        for row in doc.tables[0].rows[1:]:
            timeline.append([c.text.strip() for c in row.cells])
    return paragraphs_to_document(paragraphs, timeline=timeline)


def load_pdf(path: Path) -> dict[str, Any]:
    parts: list[str] = []
    with pdfplumber.open(str(path)) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                parts.append(text)
    return paragraphs_to_document(_split_paragraphs("\n\n".join(parts)))


def load_plain_text(text: str) -> dict[str, Any]:
    return paragraphs_to_document(_split_paragraphs(text))


def load_source(
    *,
    path: Path | None = None,
    text: str | None = None,
    title: str | None = None,
    subtitle: str | None = None,
) -> dict[str, Any]:
    if path is not None:
        suffix = path.suffix.lower()
        if suffix == ".docx":
            doc = load_docx(path)
        elif suffix == ".pdf":
            doc = load_pdf(path)
        else:
            raise ValueError(f"Unsupported file type: {suffix}")
    elif text and text.strip():
        doc = load_plain_text(text.strip())
    else:
        raise ValueError("Provide a file path or plain text content.")

    if title:
        doc["title"] = title.strip()
    if subtitle:
        doc["subtitle"] = subtitle.strip()
    return doc


def default_sample_path() -> Path | None:
    """Optional sample docx path for demos (not required)."""
    return None
